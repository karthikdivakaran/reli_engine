
import json
import pandas as pd
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
from PyQt5.QtWidgets import QMessageBox
from docx import Document
from bs4 import BeautifulSoup


def get_comp_config():
    with open("component_configuration.json", "r", encoding="utf-8") as file:
        data = json.load(file)
    return data["components"]

def get_details_values(details):
    equation_values = "λ = "
    for item in details["values"]:
        if equation_values != "λ = ":
            equation_values += " * "
        equation_values += f"{details['values'][item]['value']}"
    try:
        result = f"λ = {eval(equation_values.split('λ = ')[-1].strip())}"
    except:
        result = ""
    return result, equation_values


def export_html(details):
    env = Environment(loader=FileSystemLoader('.'))
    template = env.get_template('template/pdf_export.html')
    result, equation_val = get_details_values(details)
    html_content = template.render(details=details, result=result, equation=equation_val)

    # Save HTML file (optional)
    with open("output.html", "w") as f:
        f.write(html_content)
    # Convert HTML to PDF
    HTML(string=html_content).write_pdf("output.pdf")
    export_to_word(html_content)
    export_excel(details)


def export_to_word(html_content):
    # Create a new Word Document
    document = Document()

    # Parse HTML content
    soup = BeautifulSoup(html_content, "html.parser")

    # Simple conversion: paragraph by paragraph
    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = element.get_text(strip=True)
        if not text:
            continue
        if element.name == "h1":
            document.add_heading(text, level=1)
        elif element.name == "h2":
            document.add_heading(text, level=2)
        elif element.name == "h3":
            document.add_heading(text, level=3)
        elif element.name == "li":
            document.add_paragraph(text, style='ListBullet')
        else:
            document.add_paragraph(text)

    # Save Word file
    document.save("output.docx")

def export_excel(details, filename="output.xlsx"):
    df = pd.DataFrame(details["values"])
    df.to_excel(filename, index=False)


def confirm_delete(current_comp, title, message, comp_box):
    """ Show confirmation dialog before deleting a row """
    msg = QMessageBox(current_comp)  # Ensure it has a parent
    msg.setIcon(QMessageBox.Warning)
    msg.setWindowTitle(f"{title}")
    msg.setText(f"{message}")
    msg.setStandardButtons(QMessageBox.Ok)
    msg.setDefaultButton(QMessageBox.Ok)

    result = msg.exec_()  # Show dialog

    if result == QMessageBox.Ok:
        msg.close()  # This is redundant as QMessageBox closes automatically


def load_stylesheet(path):
    with open(path, "r") as f:
        return f.read()